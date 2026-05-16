"""Coverage for app.ingest.extract_sections — guards against the silent
content-loss bugs we hit in production (DOCX tables / headers / footers
dropped; HTML alt / meta description / display:none not handled).

We build fixtures programmatically with python-docx so the tests don't
depend on any binary asset, and write small HTML strings inline for the
HTML cases.
"""
from pathlib import Path

import pytest
from docx import Document

from app.ingest import extract_sections


# -------------------- DOCX --------------------

def _make_docx(tmp_path: Path) -> Path:
    """A docx that exercises every silent-loss path: paragraphs interleaved
    with tables, a nested table, custom header + footer text."""
    doc = Document()

    doc.add_paragraph("Intro paragraph one.")
    doc.add_paragraph("Intro paragraph two.")

    # Simple table.
    t1 = doc.add_table(rows=2, cols=3)
    t1.cell(0, 0).text = "Use Case"
    t1.cell(0, 1).text = "Customer"
    t1.cell(0, 2).text = "ABC Precision"
    t1.cell(1, 0).text = "Industry"
    t1.cell(1, 1).text = "Manufacturing"
    t1.cell(1, 2).text = "Taiwan"

    doc.add_paragraph("Between paragraph that follows the first table.")

    # Table with a nested table inside cell(0, 1).
    t2 = doc.add_table(rows=1, cols=2)
    t2.cell(0, 0).text = "Outer cell A"
    nested = t2.cell(0, 1).add_table(rows=1, cols=2)
    nested.cell(0, 0).text = "NESTED-LEFT"
    nested.cell(0, 1).text = "NESTED-RIGHT"

    doc.add_paragraph("Closing paragraph.")

    # Header + footer on the default section.
    doc.sections[0].header.add_paragraph("DOC HEADER LINE")
    doc.sections[0].footer.add_paragraph("DOC FOOTER LINE")

    path = tmp_path / "fixture.docx"
    doc.save(str(path))
    return path


def test_docx_extracts_paragraphs_and_table_cells(tmp_path):
    """The legacy bug: table cells silently dropped. Regression check."""
    sections = extract_sections(_make_docx(tmp_path))
    body = next(text for label, text in sections if label == "document")
    # Paragraph content present.
    assert "Intro paragraph one." in body
    assert "Closing paragraph." in body
    # Cells from the first (non-nested) table.
    assert "Manufacturing" in body
    assert "ABC Precision" in body


def test_docx_preserves_body_order_paragraphs_tables_interleaved(tmp_path):
    """Body order matters: the paragraph 'between' the two tables must
    appear AFTER the first table and BEFORE the nested-table block."""
    sections = extract_sections(_make_docx(tmp_path))
    body = next(text for label, text in sections if label == "document")
    pos_intro = body.index("Intro paragraph one.")
    pos_t1 = body.index("ABC Precision")
    pos_between = body.index("Between paragraph")
    pos_t2 = body.index("Outer cell A")
    pos_closing = body.index("Closing paragraph.")
    assert pos_intro < pos_t1 < pos_between < pos_t2 < pos_closing


def test_docx_recurses_into_nested_tables(tmp_path):
    """A table inside a cell should still surface its cell text."""
    sections = extract_sections(_make_docx(tmp_path))
    body = next(text for label, text in sections if label == "document")
    assert "NESTED-LEFT" in body
    assert "NESTED-RIGHT" in body


def test_docx_extracts_header_and_footer_sections(tmp_path):
    """Headers / footers used to be invisible because they're not in
    doc.element.body — they live on each section.header / .footer."""
    sections = extract_sections(_make_docx(tmp_path))
    by_label = dict(sections)
    assert by_label.get("header", "").strip() == "DOC HEADER LINE"
    assert by_label.get("footer", "").strip() == "DOC FOOTER LINE"


def test_docx_renders_table_row_separators(tmp_path):
    """Cells in a row should join with ' | ' so the chunker can keep
    row structure recognisable for downstream retrieval."""
    sections = extract_sections(_make_docx(tmp_path))
    body = next(text for label, text in sections if label == "document")
    assert "Industry | Manufacturing | Taiwan" in body


# -------------------- HTML --------------------

def _write_html(tmp_path: Path, body: str) -> Path:
    path = tmp_path / "page.html"
    path.write_text(body, encoding="utf-8")
    return path


def test_html_extracts_meta_description(tmp_path):
    path = _write_html(tmp_path, """
        <html><head>
            <meta name="description" content="Quarterly earnings report.">
        </head><body><p>Body text only.</p></body></html>
    """)
    sections = extract_sections(path)
    text = sections[0][1]
    assert "Quarterly earnings report." in text


def test_html_extracts_image_alt_and_link_title(tmp_path):
    path = _write_html(tmp_path, """
        <html><body>
            <p>See diagram below.</p>
            <img src="x.png" alt="Architecture diagram showing the API gateway">
            <a href="/spec" title="Full specification document">spec</a>
        </body></html>
    """)
    sections = extract_sections(path)
    text = sections[0][1]
    assert "Architecture diagram showing the API gateway" in text
    assert "Full specification document" in text


def test_html_strips_script_style_and_hidden_elements(tmp_path):
    path = _write_html(tmp_path, """
        <html><body>
            <script>var SECRET = "should not appear";</script>
            <style>.foo { color: red; }</style>
            <noscript>JS-disabled fallback (should be removed)</noscript>
            <div hidden>HIDDEN_ATTR_CONTENT</div>
            <div style="display: none">HIDDEN_STYLE_CONTENT</div>
            <p>VISIBLE_CONTENT</p>
        </body></html>
    """)
    text = extract_sections(path)[0][1]
    assert "VISIBLE_CONTENT" in text
    assert "SECRET" not in text
    assert "JS-disabled fallback" not in text
    assert "HIDDEN_ATTR_CONTENT" not in text
    assert "HIDDEN_STYLE_CONTENT" not in text


def test_html_keeps_visible_table_content(tmp_path):
    """get_text() walks tables — we want to make sure adding our noise
    filters didn't break the normal happy path."""
    path = _write_html(tmp_path, """
        <html><body>
            <table>
                <tr><th>Country</th><th>Revenue</th></tr>
                <tr><td>Taiwan</td><td>123</td></tr>
            </table>
        </body></html>
    """)
    text = extract_sections(path)[0][1]
    assert "Country" in text
    assert "Taiwan" in text
    assert "123" in text
